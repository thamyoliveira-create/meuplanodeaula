# -*- coding: utf-8 -*-
"""
Flask Web Application for Lesson Plan Generator
Entrypoint for Vercel Serverless Function: api/index.py
"""

import io
import os
import csv
import pathlib
import sys
from typing import Union, List, Dict, Any

from flask import Flask, request, send_file, jsonify, render_template_string
from docx import Document

BASE_DIR = pathlib.Path(__file__).resolve().parent.parent
if str(BASE_DIR) not in sys.path:
    sys.path.insert(0, str(BASE_DIR))

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16 MB limit


# --- File Path Resolvers ---

def get_default_schedule_path() -> pathlib.Path:
    candidates = [
        pathlib.Path(__file__).parent / "schedule.csv",
        BASE_DIR / "data" / "schedule.csv",
        BASE_DIR / "api" / "schedule.csv",
        BASE_DIR / "schedule.csv"
    ]
    for p in candidates:
        if p.exists():
            return p
    return candidates[0]


def get_default_template_path() -> pathlib.Path:
    candidates = [
        pathlib.Path(__file__).parent / "template.docx",
        BASE_DIR / "template" / "template.docx",
        BASE_DIR / "lesson_plan_generator" / "template" / "template.docx",
        BASE_DIR / "api" / "template.docx"
    ]
    for p in candidates:
        if p.exists():
            return p
    return candidates[0]


# --- Data Processing Helpers ---

def parse_schedule_rows(source: Union[str, pathlib.Path, io.BytesIO], filename: str = "") -> List[Dict[str, str]]:
    """Parse CSV or Excel file and return list of dict rows."""
    if isinstance(source, (str, pathlib.Path)):
        with open(str(source), "r", encoding="utf-8", errors="replace") as f:
            reader = csv.DictReader(f)
            return [dict(row) for row in reader]
    else:
        # In-memory BytesIO
        if filename.lower().endswith((".xlsx", ".xls")):
            import pandas as pd
            df = pd.read_excel(source)
            return df.to_dict(orient="records")
        else:
            text = source.read().decode("utf-8", errors="replace")
            # detect delimiter
            delimiter = ";" if ";" in text.split("\n")[0] else ","
            reader = csv.DictReader(io.StringIO(text), delimiter=delimiter)
            return [dict(row) for row in reader]


def get_schedule_metadata(rows: List[Dict[str, str]]) -> Dict[str, Any]:
    """Extract disciplines and available weeks."""
    disciplines = set()
    discipline_weeks = {}
    all_weeks = set()

    for r in rows:
        disc = r.get("Nome do componente", "").strip()
        week = r.get("Semana", "").strip()
        if disc:
            disciplines.add(disc)
            if disc not in discipline_weeks:
                discipline_weeks[disc] = set()
            if week:
                discipline_weeks[disc].add(week)
        if week:
            all_weeks.add(week)

    # Sort weeks numerically
    def sort_weeks(w_list):
        try:
            return sorted(list(w_list), key=lambda x: int(x) if x.isdigit() else 999)
        except Exception:
            return sorted(list(w_list))

    sorted_disciplines = sorted(list(disciplines))
    sorted_disc_weeks = {d: sort_weeks(discipline_weeks[d]) for d in sorted_disciplines}
    sorted_all_weeks = sort_weeks(all_weeks)

    return {
        "disciplines": sorted_disciplines,
        "discipline_weeks": sorted_disc_weeks,
        "all_weeks": sorted_all_weeks
    }


def filter_rows(rows: List[Dict[str, str]], week: str, discipline: str = "") -> List[Dict[str, str]]:
    """Filter rows matching the week and optional discipline."""
    results = []
    for r in rows:
        r_week = str(r.get("Semana", "")).strip()
        r_disc = str(r.get("Nome do componente", "")).strip()

        if r_week == str(week).strip():
            if not discipline or r_disc == discipline.strip():
                results.append(r)
    return results


def compile_lesson_data(rows: List[Dict[str, str]], teacher: str, period: str, classroom: str, discipline: str, week: str) -> Dict[str, str]:
    """Compile aggregated text fields for the lesson plan."""
    if not rows:
        # Fallback if no specific rows found
        return {
            "professor": teacher,
            "disciplina": discipline,
            "turma": classroom,
            "periodo": period,
            "semana": week,
            "habilidades": "Desenvolver competências técnicas e socioemocionais da semana.",
            "objetos": "Objetos de conhecimento previstos para a semana.",
            "conteudo": f"Aulas da Semana {week}",
            "objetivos": "Compreender e aplicar os conceitos apresentados.",
            "estrategias": "Aulas expositivas dialogadas, estudos de casos práticos e atividades em grupo.",
            "recursos": "Lousa digital, slides, computador/notebook, material impresso e internet.",
            "avaliacao": "Participação nas atividades, realização de exercícios práticos e autoavaliação.",
            "referencias": "Currículo Técnico do Estado de São Paulo e materiais didáticos oficiais."
        }

    # Extract aggregated fields from matching rows
    aulas_titles = []
    temas = set()
    hab_tecnicas = set()
    hab_socio = set()
    obj_conhecimento = set()
    objetivos = []

    for idx, r in enumerate(rows, 1):
        tema = r.get("Tema da semana", "").strip()
        if tema: temas.add(tema)

        titulo = r.get("Título da aula", "").strip()
        if titulo: aulas_titles.append(f"• {titulo}")

        ht = r.get("Habilidades técnicas", "").strip()
        if ht: hab_tecnicas.add(f"• {ht}")

        hs = r.get("Habildades socioemocionais", "").strip()
        if hs: hab_socio.add(f"• {hs}")

        oc = r.get("Objeto de conhecimento", "").strip()
        if oc: obj_conhecimento.add(f"• {oc}")

        obj = r.get("Objetivo da aula", "").strip()
        if obj: objetivos.append(f"• {obj}")

    tema_str = " | ".join(temas) if temas else f"Semana {week}"
    conteudo_str = f"Tema: {tema_str}\n" + "\n".join(aulas_titles) if aulas_titles else tema_str
    habilidades_str = "Habilidades Técnicas:\n" + "\n".join(hab_tecnicas)
    if hab_socio:
        habilidades_str += "\n\nHabilidades Socioemocionais:\n" + "\n".join(hab_socio)

    return {
        "professor": teacher,
        "disciplina": discipline or (rows[0].get("Nome do componente", "") if rows else ""),
        "turma": classroom,
        "periodo": period,
        "semana": week,
        "habilidades": habilidades_str,
        "objetos": "\n".join(obj_conhecimento) if obj_conhecimento else "Conceitos fundamentais da disciplina.",
        "conteudo": conteudo_str,
        "objetivos": "\n".join(objetivos) if objetivos else "Desenvolver os conhecimentos previstos.",
        "estrategias": "Aulas expositivas dialogadas, estudos de casos práticos, dinâmicas e resolução de problemas em grupo.",
        "recursos": "Quadro, projetor multimídia, computadores do laboratório, material de apoio impresso/digital.",
        "avaliacao": "Avaliação contínua, participação ativa nas discussões, entrega de exercícios e produções práticas.",
        "referencias": "Diretrizes Curriculares da Educação Profissional Técnica - SEDUC-SP e bibliografia recomendada."
    }


def generate_filled_docx(lesson_data: Dict[str, str], custom_template: io.BytesIO = None) -> io.BytesIO:
    """Fill the template document with lesson plan fields."""
    template_path = get_default_template_path()
    if custom_template:
        doc = Document(custom_template)
    else:
        doc = Document(str(template_path))

    # Helper mapping for generic {{placeholders}}
    placeholders = {
        "Professor": lesson_data["professor"],
        "Nome do componente": lesson_data["disciplina"],
        "Disciplina": lesson_data["disciplina"],
        "Componente curricular": lesson_data["disciplina"],
        "Ano/Série/Turma": lesson_data["turma"],
        "Turma": lesson_data["turma"],
        "Período de realização": lesson_data["periodo"],
        "Data": lesson_data["periodo"],
        "Semana": lesson_data["semana"],
        "Habilidades": lesson_data["habilidades"],
        "Objetos de Conhecimento": lesson_data["objetos"],
        "Conteúdo": lesson_data["conteudo"],
        "Objetivos": lesson_data["objetivos"],
        "Estratégias": lesson_data["estrategias"],
        "Recursos didáticos": lesson_data["recursos"],
        "Critérios de Avaliação": lesson_data["avaliacao"],
        "Referências": lesson_data["referencias"]
    }

    # Replace in all paragraphs
    for p in doc.paragraphs:
        for k, v in placeholders.items():
            token = f"{{{{{k}}}}}"
            if token in p.text:
                p.text = p.text.replace(token, v)

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # 1. Check placeholders
                for k, v in placeholders.items():
                    token = f"{{{{{k}}}}}"
                    if token in cell.text:
                        cell.text = cell.text.replace(token, v)

                # 2. Check standard template labeled rows
                c_text = cell.text.strip()
                if c_text == "Professor":
                    cell.text = f"Professor: {lesson_data['professor']}"
                elif c_text.startswith("Componente curricular:"):
                    cell.text = f"Componente curricular: {lesson_data['disciplina']}"
                elif "Ano/Série/Turma" in c_text:
                    cell.text = f"3. Ano/Série/Turma: {lesson_data['turma']}"
                elif "Período de realização" in c_text:
                    cell.text = f"4. Período de realização: {lesson_data['periodo']}"
                elif c_text.startswith("5. Habilidades") or c_text.startswith("5 . Habilidades"):
                    cell.text = f"5. Habilidades:\n{lesson_data['habilidades']}"
                elif "Objetos de Conhecimento" in c_text:
                    cell.text = f"6. Objetos de Conhecimento:\n{lesson_data['objetos']}"
                elif "Conteúdo" in c_text:
                    cell.text = f"7. Conteúdo:\n{lesson_data['conteudo']}"
                elif "Objetivos" in c_text and not "Objetos" in c_text:
                    cell.text = f"8. Objetivos:\n{lesson_data['objetivos']}"
                elif "Estratégias" in c_text:
                    cell.text = f"9. Estratégias:\n{lesson_data['estrategias']}"
                elif "Recursos didáticos" in c_text:
                    cell.text = f"10. Recursos didáticos:\n{lesson_data['recursos']}"
                elif "Critérios de Avaliação" in c_text:
                    cell.text = f"11. Critérios de Avaliação:\n{lesson_data['avaliacao']}"
                elif "Referências" in c_text:
                    cell.text = f"12. Referências:\n{lesson_data['referencias']}"

    out_stream = io.BytesIO()
    doc.save(out_stream)
    out_stream.seek(0)
    return out_stream


# --- Embedded HTML Template ---

HTML_TEMPLATE = """<!DOCTYPE html>
<html lang="pt-BR">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Gerador de Plano de Aula</title>
    <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.3/dist/css/bootstrap.min.css" rel="stylesheet">
    <link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/bootstrap-icons@1.11.3/font/bootstrap-icons.min.css">
    <style>
        :root {
            --primary-gradient: linear-gradient(135deg, #4f46e5 0%, #7c3aed 100%);
            --card-shadow: 0 10px 25px -5px rgba(0, 0, 0, 0.07), 0 8px 10px -6px rgba(0, 0, 0, 0.04);
        }
        body {
            background-color: #f1f5f9;
            font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, "Helvetica Neue", Arial, sans-serif;
            color: #1e293b;
            min-height: 100vh;
        }
        .navbar-custom {
            background: var(--primary-gradient);
            padding: 1rem 0;
            box-shadow: 0 4px 15px rgba(79, 70, 229, 0.25);
        }
        .main-card {
            background: white;
            border-radius: 1.25rem;
            box-shadow: var(--card-shadow);
            padding: 2.2rem;
            border: 1px solid #e2e8f0;
        }
        .btn-gradient {
            background: var(--primary-gradient);
            border: none;
            color: white;
            padding: 0.75rem 1.5rem;
            font-weight: 600;
            border-radius: 0.75rem;
            transition: all 0.2s ease;
        }
        .btn-gradient:hover {
            transform: translateY(-1px);
            box-shadow: 0 6px 20px rgba(79, 70, 229, 0.35);
            color: white;
        }
        .form-label {
            font-weight: 600;
            color: #334155;
            font-size: 0.92rem;
            margin-bottom: 0.4rem;
        }
        .form-control, .form-select {
            border-radius: 0.65rem;
            border: 1.5px solid #cbd5e1;
            padding: 0.65rem 0.9rem;
            font-size: 0.95rem;
        }
        .form-control:focus, .form-select:focus {
            border-color: #6366f1;
            box-shadow: 0 0 0 4px rgba(99, 102, 241, 0.15);
        }
        .preview-box {
            background: #ffffff;
            border-radius: 1rem;
            border: 1px solid #e2e8f0;
            padding: 2rem;
            box-shadow: var(--card-shadow);
        }
        .field-card {
            background: #f8fafc;
            border: 1px solid #e2e8f0;
            border-left: 4px solid #4f46e5;
            border-radius: 0.5rem;
            padding: 1rem 1.2rem;
            margin-bottom: 1rem;
        }
        .field-title {
            font-weight: 700;
            font-size: 0.8rem;
            text-transform: uppercase;
            letter-spacing: 0.05em;
            color: #64748b;
            margin-bottom: 0.3rem;
        }
        .field-text {
            color: #0f172a;
            font-size: 0.95rem;
            white-space: pre-line;
            line-height: 1.5;
        }
        .school-header {
            text-align: center;
            border-bottom: 2px solid #e2e8f0;
            padding-bottom: 1rem;
            margin-bottom: 1.5rem;
        }
        .school-header h6 {
            font-weight: 700;
            margin-bottom: 0.2rem;
            font-size: 0.85rem;
            color: #475569;
        }
        @media print {
            .no-print { display: none !important; }
            body { background: white !important; padding: 0 !important; }
            .preview-box { border: none !important; box-shadow: none !important; padding: 0 !important; }
            .field-card { border: 1px solid #ddd !important; border-left: 4px solid #333 !important; }
        }
    </style>
</head>
<body>

    <nav class="navbar navbar-dark navbar-custom no-print">
        <div class="container">
            <a class="navbar-brand d-flex align-items-center gap-2 fw-bold fs-4" href="#">
                <i class="bi bi-journal-text fs-3"></i>
                Gerador de Plano de Aula
            </a>
            <span class="navbar-text text-white-50 d-none d-md-inline">
                EE. Prof. Dinora Marcondes Gomes &bull; Cursos Técnicos
            </span>
        </div>
    </nav>

    <div class="container my-4">
        <div class="main-card mb-4 no-print">
            <div class="d-flex align-items-center justify-content-between mb-3 pb-2 border-bottom">
                <div>
                    <h3 class="fw-bold mb-1 text-dark">Informações do Plano de Aula</h3>
                    <p class="text-muted mb-0 small">Preencha os dados da aula. O conteúdo curricular e as habilidades já estão integrados.</p>
                </div>
                <span class="badge bg-success-subtle text-success border border-success-subtle px-3 py-2 rounded-pill">
                    <i class="bi bi-check-circle-fill me-1"></i> Conteúdo Carregado
                </span>
            </div>

            <form id="planForm" action="/generate" method="POST" enctype="multipart/form-data">
                <div class="row g-3">
                    <div class="col-md-6">
                        <label for="teacher" class="form-label">
                            <i class="bi bi-person-badge text-primary me-1"></i> Nome do(a) Professor(a)
                        </label>
                        <input type="text" class="form-control" id="teacher" name="teacher" placeholder="Ex: Prof. Thamy Oliveira" required>
                    </div>

                    <div class="col-md-6">
                        <label for="period" class="form-label">
                            <i class="bi bi-calendar-range text-primary me-1"></i> Período de Realização / Data
                        </label>
                        <input type="text" class="form-control" id="period" name="period" placeholder="Ex: Semana de 24 a 28 de Fevereiro" required>
                    </div>

                    <div class="col-md-4">
                        <label for="classroom" class="form-label">
                            <i class="bi bi-people text-primary me-1"></i> Ano / Série / Turma
                        </label>
                        <input type="text" class="form-control" id="classroom" name="classroom" placeholder="Ex: 1º MTEC - Administração" required>
                    </div>

                    <div class="col-md-5">
                        <label for="discipline" class="form-label">
                            <i class="bi bi-book text-primary me-1"></i> Disciplina (Componente Curricular)
                        </label>
                        <select class="form-select" id="discipline" name="discipline" onchange="updateWeeksDropdown()">
                        </select>
                    </div>

                    <div class="col-md-3">
                        <label for="week" class="form-label">
                            <i class="bi bi-calendar3-week text-primary me-1"></i> Semana do Curso
                        </label>
                        <select class="form-select" id="week" name="week" required>
                        </select>
                    </div>
                </div>

                <div class="accordion accordion-flush mt-3" id="advancedAccordion">
                    <div class="accordion-item border-0">
                        <h2 class="accordion-header">
                            <button class="accordion-button collapsed px-0 py-2 text-secondary small bg-transparent shadow-none" type="button" data-bs-toggle="collapse" data-bs-target="#advancedFiles">
                                <i class="bi bi-sliders me-1"></i> Opções Avançadas: Usar outra planilha ou modelo personalizado
                            </button>
                        </h2>
                        <div id="advancedFiles" class="accordion-collapse collapse">
                            <div class="p-3 bg-light rounded-3 mt-2 border">
                                <div class="row g-3">
                                    <div class="col-md-6">
                                        <label class="form-label small fw-bold">Substituir Planilha (.xlsx ou .csv)</label>
                                        <input type="file" class="form-control form-control-sm" id="customSpreadsheet" name="custom_spreadsheet" accept=".xlsx, .xls, .csv">
                                    </div>
                                    <div class="col-md-6">
                                        <label class="form-label small fw-bold">Substituir Modelo Word (.docx)</label>
                                        <input type="file" class="form-control form-control-sm" id="customTemplate" name="custom_template" accept=".docx">
                                    </div>
                                </div>
                            </div>
                        </div>
                    </div>
                </div>

                <div class="d-flex flex-wrap gap-2 mt-4 pt-2 border-top">
                    <button type="submit" class="btn btn-gradient px-4 py-2" id="btnDownload">
                        <i class="bi bi-file-earmark-word me-1 fs-5 align-middle"></i> Baixar Word (.docx)
                    </button>
                    <button type="button" class="btn btn-outline-primary px-4 py-2" id="btnPreview">
                        <i class="bi bi-eye me-1 align-middle"></i> Pré-visualizar na Tela
                    </button>
                </div>
            </form>
        </div>

        <div id="alertBox" class="alert d-none no-print" role="alert"></div>

        <div id="previewArea" class="preview-box d-none">
            <div class="d-flex justify-content-between align-items-center mb-3 pb-2 border-bottom no-print">
                <h5 class="fw-bold mb-0 text-primary">
                    <i class="bi bi-file-earmark-check me-1"></i> Visualização do Plano de Aula
                </h5>
                <button class="btn btn-sm btn-dark" onclick="window.print()">
                    <i class="bi bi-printer me-1"></i> Imprimir / Salvar em PDF
                </button>
            </div>

            <div class="school-header">
                <h6>GOVERNO DO ESTADO DE SÃO PAULO – SECRETARIA DA EDUCAÇÃO</h6>
                <h6>DIRETORIA DE ENSINO - REGIÃO DE ARARAQUARA</h6>
                <h6>EE. PROF. DINORA MARCONDES GOMES</h6>
                <p class="text-muted small mb-0">RUA EMILIA GALLI, 549 - CENTRO - AMÉRICO BRASILIENSE - SP | TEL: (16) 3392-1335</p>
                <h5 class="fw-bold mt-2 text-dark">PLANO DE AULA - 2026</h5>
            </div>

            <div class="row g-2 mb-3">
                <div class="col-md-6">
                    <div class="p-2 border rounded bg-light">
                        <strong>Professor(a):</strong> <span id="prevTeacher">-</span>
                    </div>
                </div>
                <div class="col-md-6">
                    <div class="p-2 border rounded bg-light">
                        <strong>Componente Curricular:</strong> <span id="prevDiscipline">-</span>
                    </div>
                </div>
                <div class="col-md-6">
                    <div class="p-2 border rounded bg-light">
                        <strong>Ano / Série / Turma:</strong> <span id="prevClassroom">-</span>
                    </div>
                </div>
                <div class="col-md-6">
                    <div class="p-2 border rounded bg-light">
                        <strong>Período de Realização:</strong> <span id="prevPeriod">-</span>
                    </div>
                </div>
            </div>

            <div id="previewFields"></div>
        </div>
    </div>

    <footer class="text-center text-muted py-4 mt-5 no-print border-top">
        <div class="container">
            <small>Gerador de Plano de Aula &bull; Desenvolvido para Cursos Técnicos</small>
        </div>
    </footer>

    <script src="https://cdn.jsdelivr.net/npm/bootstrap@5.3.3/dist/js/bootstrap.bundle.min.js"></script>

    <script>
        let courseData = {};

        async function loadInitialData() {
            try {
                const res = await fetch('/api/initial-data');
                const data = await res.json();
                if (data.success) {
                    courseData = data;
                    const discSelect = document.getElementById('discipline');
                    discSelect.innerHTML = '';
                    
                    data.disciplines.forEach(d => {
                        const opt = document.createElement('option');
                        opt.value = d;
                        opt.textContent = d;
                        discSelect.appendChild(opt);
                    });

                    updateWeeksDropdown();
                }
            } catch (err) {
                console.error("Erro ao carregar dados iniciais:", err);
            }
        }

        function updateWeeksDropdown() {
            const discipline = document.getElementById('discipline').value;
            const weekSelect = document.getElementById('week');
            weekSelect.innerHTML = '';

            const weeks = (courseData.discipline_weeks && courseData.discipline_weeks[discipline]) 
                ? courseData.discipline_weeks[discipline] 
                : (courseData.all_weeks || Array.from({length: 28}, (_, i) => String(i+1)));

            weeks.forEach(w => {
                const opt = document.createElement('option');
                opt.value = w;
                opt.textContent = `Semana ${w}`;
                weekSelect.appendChild(opt);
            });
        }

        document.getElementById('btnPreview').addEventListener('click', async function() {
            const teacher = document.getElementById('teacher').value.trim();
            const period = document.getElementById('period').value.trim();
            const classroom = document.getElementById('classroom').value.trim();
            const discipline = document.getElementById('discipline').value;
            const week = document.getElementById('week').value;

            if (!teacher || !period || !classroom || !week) {
                showAlert("Por favor, preencha o Nome do Professor, Data/Período, Turma e Semana.");
                return;
            }

            const btn = document.getElementById('btnPreview');
            btn.disabled = true;
            btn.innerHTML = `<span class="spinner-border spinner-border-sm me-1"></span> Carregando...`;

            const form = document.getElementById('planForm');
            const formData = new FormData(form);

            try {
                const res = await fetch('/api/preview-plan', { method: 'POST', body: formData });
                const result = await res.json();

                if (!result.success) {
                    showAlert(result.error || "Erro ao gerar pré-visualização.");
                } else {
                    document.getElementById('alertBox').classList.add('d-none');
                    document.getElementById('prevTeacher').textContent = teacher;
                    document.getElementById('prevDiscipline').textContent = discipline;
                    document.getElementById('prevClassroom').textContent = classroom;
                    document.getElementById('prevPeriod').textContent = period;

                    const fieldsDiv = document.getElementById('previewFields');
                    fieldsDiv.innerHTML = '';

                    const sections = [
                        { label: '5. Habilidades (Técnicas e Socioemocionais)', value: result.data.habilidades },
                        { label: '6. Objetos de Conhecimento', value: result.data.objetos },
                        { label: '7. Conteúdo das Aulas', value: result.data.conteudo },
                        { label: '8. Objetivos das Aulas', value: result.data.objetivos },
                        { label: '9. Estratégias Metodológicas', value: result.data.estrategias },
                        { label: '10. Recursos Didáticos', value: result.data.recursos },
                        { label: '11. Critérios de Avaliação', value: result.data.avaliacao },
                        { label: '12. Referências', value: result.data.referencias }
                    ];

                    sections.forEach(sec => {
                        const card = document.createElement('div');
                        card.className = 'field-card';
                        card.innerHTML = `
                            <div class="field-title">${sec.label}</div>
                            <div class="field-text">${sec.value || 'Não informado'}</div>
                        `;
                        fieldsDiv.appendChild(card);
                    });

                    const prevArea = document.getElementById('previewArea');
                    prevArea.classList.remove('d-none');
                    prevArea.scrollIntoView({ behavior: 'smooth' });
                }
            } catch (err) {
                showAlert("Erro ao conectar com o servidor.");
            } finally {
                btn.disabled = false;
                btn.innerHTML = `<i class="bi bi-eye me-1 align-middle"></i> Pré-visualizar na Tela`;
            }
        });

        function showAlert(msg, type = 'danger') {
            const alertBox = document.getElementById('alertBox');
            alertBox.className = `alert alert-${type} alert-dismissible fade show no-print`;
            alertBox.innerHTML = `<span>${msg}</span><button type="button" class="btn-close" data-bs-dismiss="alert"></button>`;
            alertBox.classList.remove('d-none');
            alertBox.scrollIntoView({ behavior: 'smooth' });
        }

        window.addEventListener('DOMContentLoaded', loadInitialData);
    </script>
</body>
</html>"""


# --- Flask Routes ---

@app.route("/", methods=["GET"])
def index():
    return render_template_string(HTML_TEMPLATE)


@app.route("/api/health", methods=["GET"])
def health():
    return jsonify({"status": "ok", "app": "Gerador de Plano de Aula"})


@app.route("/api/initial-data", methods=["GET"])
def initial_data():
    try:
        schedule_path = get_default_schedule_path()
        rows = parse_schedule_rows(schedule_path)
        meta = get_schedule_metadata(rows)
        return jsonify({"success": True, **meta})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 400


@app.route("/api/preview-plan", methods=["POST"])
def preview_plan():
    teacher = request.form.get("teacher", "").strip()
    period = request.form.get("period", "").strip()
    classroom = request.form.get("classroom", "").strip()
    discipline = request.form.get("discipline", "").strip()
    week = request.form.get("week", "").strip()

    if not week:
        return jsonify({"success": False, "error": "Semana não informada."}), 400

    try:
        # Check if custom spreadsheet was uploaded
        if "custom_spreadsheet" in request.files and request.files["custom_spreadsheet"].filename:
            f = request.files["custom_spreadsheet"]
            rows = parse_schedule_rows(io.BytesIO(f.read()), filename=f.filename)
        else:
            schedule_path = get_default_schedule_path()
            rows = parse_schedule_rows(schedule_path)

        filtered = filter_rows(rows, week, discipline)
        lesson_data = compile_lesson_data(filtered, teacher, period, classroom, discipline, week)
        return jsonify({"success": True, "data": lesson_data})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 400


@app.route("/generate", methods=["POST"])
def generate():
    teacher = request.form.get("teacher", "").strip()
    period = request.form.get("period", "").strip()
    classroom = request.form.get("classroom", "").strip()
    discipline = request.form.get("discipline", "").strip()
    week = request.form.get("week", "").strip()

    if not week:
        return "Erro: Semana não informada.", 400

    try:
        if "custom_spreadsheet" in request.files and request.files["custom_spreadsheet"].filename:
            f = request.files["custom_spreadsheet"]
            rows = parse_schedule_rows(io.BytesIO(f.read()), filename=f.filename)
        else:
            schedule_path = get_default_schedule_path()
            rows = parse_schedule_rows(schedule_path)

        custom_tpl = None
        if "custom_template" in request.files and request.files["custom_template"].filename:
            t = request.files["custom_template"]
            custom_tpl = io.BytesIO(t.read())

        filtered = filter_rows(rows, week, discipline)
        lesson_data = compile_lesson_data(filtered, teacher, period, classroom, discipline, week)
        docx_buffer = generate_filled_docx(lesson_data, custom_template=custom_tpl)

        clean_disc = "".join(c for c in discipline if c.isalnum() or c in " _-").strip().replace(" ", "_")
        filename = f"Plano_de_Aula_Semana_{week}_{clean_disc or 'Tecnico'}.docx"

        return send_file(
            docx_buffer,
            as_attachment=True,
            download_name=filename,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except Exception as e:
        return f"Erro ao gerar plano de aula: {str(e)}", 400


if __name__ == "__main__":
    app.run(debug=True, host="127.0.0.1", port=5000)
