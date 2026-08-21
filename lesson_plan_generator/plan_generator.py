# -*- coding: utf-8 -*-
"""
Core logic for generating a lesson plan (.docx and preview).
"""

import io
import pathlib
from typing import Union, List, Dict, Any
import pandas as pd
from docx import Document

DEFAULT_TEMPLATE_PATH = pathlib.Path(__file__).parent / "template" / "template.docx"


def load_spreadsheet(file_or_path: Union[str, pathlib.Path, io.BytesIO], filename: str = "") -> pd.DataFrame:
    """Load spreadsheet from path or BytesIO."""
    if isinstance(file_or_path, (str, pathlib.Path)):
        path = pathlib.Path(file_or_path)
        if not path.exists():
            raise FileNotFoundError(f"Planilha não encontrada: {path}")
        if path.suffix.lower() in {".xlsx", ".xls"}:
            df = pd.read_excel(path)
        else:
            df = pd.read_csv(path, delimiter=None, engine="python")
    else:
        # In-memory BytesIO
        if filename.lower().endswith((".xlsx", ".xls")):
            df = pd.read_excel(file_or_path)
        else:
            try:
                df = pd.read_csv(file_or_path, delimiter=",")
            except Exception:
                file_or_path.seek(0)
                df = pd.read_csv(file_or_path, delimiter=";")

    # Strip column names
    df.columns = [str(col).strip() for col in df.columns]
    return df


def find_week_column(df: pd.DataFrame) -> str:
    """Find the column name corresponding to week."""
    for col in df.columns:
        if "semana" in col.lower():
            return col
    raise KeyError("Não foi encontrada a coluna de 'Semana' na planilha enviada.")


def get_available_weeks(df: pd.DataFrame) -> List[str]:
    """Get list of unique weeks present in the spreadsheet."""
    week_col = find_week_column(df)
    weeks = df[week_col].dropna().astype(str).str.strip().unique().tolist()
    # Sort nicely if numeric
    try:
        weeks.sort(key=lambda x: int(x) if x.isdigit() else x)
    except Exception:
        pass
    return [w for w in weeks if w]


def filter_by_week(df: pd.DataFrame, week: str) -> pd.DataFrame:
    """Filter rows matching the selected week."""
    week_col = find_week_column(df)
    filtered = df[df[week_col].astype(str).str.strip() == str(week).strip()]
    if filtered.empty:
        raise ValueError(f"Nenhum registro encontrado para a semana '{week}'.")
    return filtered


def _replace_in_paragraph(paragraph, placeholder_map: Dict[str, str]):
    """
    Robust placeholder replacement across runs in a paragraph.
    Handles {{ColumnName}} even if split across runs by Word.
    """
    full_text = "".join(run.text for run in paragraph.runs)
    has_match = any(f"{{{{{key}}}}}" in full_text for key in placeholder_map)

    if not has_match:
        return

    # First attempt simple run-by-run replacement
    for run in paragraph.runs:
        for key, val in placeholder_map.items():
            token = f"{{{{{key}}}}}"
            if token in run.text:
                run.text = run.text.replace(token, val)

    # Re-check if any tokens are still present across run boundaries
    full_text = "".join(run.text for run in paragraph.runs)
    if any(f"{{{{{key}}}}}" in full_text for key in placeholder_map):
        for key, val in placeholder_map.items():
            token = f"{{{{{key}}}}}"
            full_text = full_text.replace(token, val)
        # Put replaced text in first run and clear others
        if paragraph.runs:
            paragraph.runs[0].text = full_text
            for run in paragraph.runs[1:]:
                run.text = ""


def render_docx(
    df: pd.DataFrame,
    template_source: Union[str, pathlib.Path, io.BytesIO] = None
) -> io.BytesIO:
    """Render lesson plan into DOCX and return BytesIO buffer."""
    if template_source is None:
        template_source = DEFAULT_TEMPLATE_PATH

    if isinstance(template_source, (str, pathlib.Path)):
        template_path = pathlib.Path(template_source)
        if not template_path.exists():
            raise FileNotFoundError(f"Arquivo de modelo não encontrado: {template_path}")
        doc = Document(str(template_path))
    else:
        doc = Document(template_source)

    # Build placeholder mapping
    placeholder_map = {}
    for col in df.columns:
        values = df[col].dropna().astype(str).tolist()
        placeholder_map[col] = "\n".join(values)

    # Process all paragraphs
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, placeholder_map)

    # Process all tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, placeholder_map)

    out_stream = io.BytesIO()
    doc.save(out_stream)
    out_stream.seek(0)
    return out_stream


def get_plan_preview(df: pd.DataFrame) -> Dict[str, Any]:
    """Get structured data dictionary for previewing in UI."""
    data = {}
    for col in df.columns:
        values = df[col].dropna().astype(str).tolist()
        data[col] = values
    return data
