# -*- coding: utf-8 -*-
"""Core logic for generating a lesson plan.

The function ``generate_lesson_plan`` receives:
- ``spreadsheet_path`` – path to a CSV (or .xlsx) with the schedule.
- ``week`` – identifier used to filter rows (exact string match on the column named "Semana").
- ``output_pdf_path`` – where the final PDF will be written.

Steps performed:
1. Load the spreadsheet with pandas (supports CSV and Excel).
2. Keep only rows where the ``Semana`` column matches ``week``.
3. Load the .docx template (stored at ``lesson_plan_generator/template/template.docx``).
4. Replace placeholders inside the document.  Placeholders are written in the template as
   ``{{column_name}}`` (e.g. ``{{Título da aula}}``).  For each column present in the filtered
   rows we replace the placeholder with the cell value.  If a placeholder appears multiple
   times it is replaced everywhere.
5. Save the filled document as a temporary ``filled.docx``.
6. Convert the temporary DOCX to PDF using ``docx2pdf`` (available on macOS and Windows).
   The conversion is performed only when the ``docx2pdf`` package is importable; otherwise a
   clear ``RuntimeError`` is raised.
"""

import pathlib
import sys
from typing import Union

import pandas as pd
from docx import Document

# Optional import – conversion to PDF
try:
    from docx2pdf import convert
except Exception:  # pragma: no cover
    convert = None

TEMPLATE_PATH = pathlib.Path(__file__).parent / "template" / "template.docx"


def _load_spreadsheet(path: Union[str, pathlib.Path]) -> pd.DataFrame:
    path = pathlib.Path(path)
    if not path.exists():
        raise FileNotFoundError(f"Spreadsheet not found: {path}")
    if path.suffix.lower() in {".xlsx", ".xls"}:
        df = pd.read_excel(path)
    else:
        df = pd.read_csv(path, delimiter=",")
    # Normalise column names – strip whitespace and keep original case (used in placeholders)
    df.columns = [col.strip() for col in df.columns]
    return df


def _filter_week(df: pd.DataFrame, week: str) -> pd.DataFrame:
    # Try to locate a column that represents the week identifier.
    # Common names in the provided sheet are "Semana" or "Semana ".
    week_col = None
    for col in df.columns:
        if "semana" in col.lower():
            week_col = col
            break
    if week_col is None:
        raise KeyError("Could not find a column representing the week (e.g., 'Semana') in the spreadsheet.")
    filtered = df[df[week_col].astype(str).str.strip() == str(week).strip()]
    if filtered.empty:
        raise ValueError(f"No rows found for week '{week}'. Check the week identifier and spreadsheet content.")
    return filtered


def _render_template(df: pd.DataFrame) -> pathlib.Path:
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Template file not found at {TEMPLATE_PATH}")
    doc = Document(str(TEMPLATE_PATH))
    # For each placeholder {{column_name}} replace with the first value from the dataframe.
    # The sheet may contain multiple rows for the same week (e.g., several lessons). In that
    # case we concatenate the values with a line break.
    placeholder_map = {}
    for col in df.columns:
        # Concatenate all non‑null values for this column, separated by newlines.
        values = df[col].dropna().astype(str).tolist()
        placeholder_map[col] = "\n".join(values)
    # Helper to replace text in a run
    def replace_in_run(run, mapping):
        for key, val in mapping.items():
            token = f"{{{{{key}}}}}"  # {{column}}
            if token in run.text:
                run.text = run.text.replace(token, val)
    # Iterate over paragraphs and tables
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            replace_in_run(run, placeholder_map)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        replace_in_run(run, placeholder_map)
    # Save the filled document to a temporary location
    filled_path = pathlib.Path(__file__).parent / "filled.docx"
    doc.save(str(filled_path))
    return filled_path


def _convert_to_pdf(docx_path: pathlib.Path, pdf_path: pathlib.Path) -> None:
    if convert is None:
        raise RuntimeError(
            "docx2pdf is not available. Install it (pip install docx2pdf) and ensure you have a "
            "compatible OS (macOS or Windows) to convert DOCX to PDF."
        )
    # docx2pdf expects string paths
    convert(str(docx_path), str(pdf_path))


def generate_lesson_plan(
    spreadsheet_path: Union[str, pathlib.Path],
    week: str,
    output_pdf_path: Union[str, pathlib.Path],
) -> None:
    """Generate a lesson‑plan PDF for the given week.

    Args:
        spreadsheet_path: Path to the CSV/Excel file that holds the schedule.
        week: Week identifier (exact match on the ``Semana`` column).
        output_pdf_path: Destination PDF file.
    """
    df = _load_spreadsheet(spreadsheet_path)
    filtered = _filter_week(df, week)
    filled_docx = _render_template(filtered)
    pdf_path = pathlib.Path(output_pdf_path)
    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    _convert_to_pdf(filled_docx, pdf_path)
    # Cleanup temporary filled.docx
    try:
        filled_docx.unlink()
    except Exception:
        pass

if __name__ == "__main__":  # pragma: no cover
    # Simple manual test when the script is executed directly.
    if len(sys.argv) != 4:
        print("Usage: python -m lesson_plan_generator.plan_generator <spreadsheet> <week> <output.pdf>")
        sys.exit(1)
    generate_lesson_plan(sys.argv[1], sys.argv[2], sys.argv[3])
